from __future__ import annotations

import json
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document


BASE_DIR = Path(__file__).resolve().parents[1]
DATA_MODEL_DIR = BASE_DIR / "data_model"
OUT_DIR = DATA_MODEL_DIR / "generated"


DOC_SPECS = {
    "unified": {
        "path": DATA_MODEL_DIR / "Unified_subcategories_mapping_v4_detailed.docx",
        "category": None,
    },
    "document": {
        "path": DATA_MODEL_DIR / "Document_types_revised_v4.docx",
        "category": "Document",
    },
    "image": {
        "path": DATA_MODEL_DIR / "Image_types_and_features_v4.docx",
        "category": "Image",
    },
    "audio": {
        "path": DATA_MODEL_DIR / "Audio_types_and_features_v4.docx",
        "category": "Audio",
    },
    "video": {
        "path": DATA_MODEL_DIR / "Video_types_and_features_v4.docx",
        "category": "Video",
    },
    "dataset": {
        "path": DATA_MODEL_DIR / "Dataset_types_and_features_v4.docx",
        "category": "Dataset",
    },
    "software": {
        "path": DATA_MODEL_DIR / "Software_types_and_features_v4.docx",
        "category": "Software",
    },
}


UNIFIED_APPLICABILITY_HINTS = {
    "how_to_guides": ["Document", "Audio", "Video"],
    "explainers": ["Document", "Audio", "Video"],
    "technical_and_research_content": ["Document", "Dataset"],
    "case_studies": ["Document", "Audio", "Video"],
    "field_demonstrations": ["Video", "Image"],
    "interviews": ["Audio", "Video"],
    "panel_discussions": ["Audio", "Video"],
    "q_and_a_sessions": ["Audio", "Video"],
    "talks_and_lectures": ["Document", "Audio", "Video"],
    "tool_walkthroughs": ["Video", "Software"],
    "simulations": ["Video", "Software"],
    "charts_graphs": ["Image"],
    "maps": ["Image", "Dataset", "Software"],
    "infographics": ["Image"],
    "diagrams": ["Image", "Document"],
    "photos": ["Image"],
    "diagnostic_images": ["Image"],
    "datasets": ["Dataset"],
    "software_tools": ["Software"],
    "monitoring_data": ["Dataset"],
    "input_data": ["Dataset"],
    "output_data": ["Dataset"],
    "survey_data": ["Dataset"],
    "templates": ["Document"],
}


CATEGORY_TYPE_ALIASES = {
    "Image": {
        "chartgraph": "chartgraph",
        "infographic": "infographic",
        "diagramschematic": "diagramschematic",
        "diagram": "diagramschematic",
        "map": "map",
        "fieldphoto": "fieldobservationalphotograph",
        "fieldobservationalphotograph": "fieldobservationalphotograph",
        "diagnosticphoto": "diagnosticphotograph",
        "diagnosticphotograph": "diagnosticphotograph",
        "equipmentphoto": "equipmentinfrastructurephotograph",
        "equipmentphotograph": "equipmentinfrastructurephotograph",
        "equipmentinfrastructurephotograph": "equipmentinfrastructurephotograph",
        "aerialremotesensingimage": "aerialremotesensingimage",
    },
    "Audio": {
        "interview": "interviewpractitionerperspective",
        "interviewpractitionerperspective": "interviewpractitionerperspective",
        "explainer": "expertcommentaryexplainer",
        "expertcommentaryexplainer": "expertcommentaryexplainer",
        "casestudy": "casestudypracticestory",
        "casestudypracticestory": "casestudypracticestory",
        "panel": "paneldiscussion",
        "paneldiscussion": "paneldiscussion",
        "qa": "qa",
        "lecture": "talklecture",
        "talklecture": "talklecture",
        "howto": "howtoprocedureguide",
        "howtoprocedureguide": "howtoprocedureguide",
    },
    "Video": {
        "walkthrough": "fielddemonstrationwalkthrough",
        "fielddemonstrationwalkthrough": "fielddemonstrationwalkthrough",
        "howto": "howtoproceduredemonstration",
        "howtoproceduredemonstration": "howtoproceduredemonstration",
        "casestudy": "casestudy",
        "explainer": "explainerdocumentary",
        "explainerdocumentary": "explainerdocumentary",
        "interview": "interviewpractitionerperspective",
        "interviewpractitionerperspective": "interviewpractitionerperspective",
        "qa": "expertqasession",
        "expertqasession": "expertqasession",
        "panel": "paneldiscussion",
        "paneldiscussion": "paneldiscussion",
        "presentation": "recordedpresentationwebinar",
        "recordedpresentationwebinar": "recordedpresentationwebinar",
        "toolwalkthrough": "toolmachinerysoftwarewalkthrough",
        "toolmachinerysoftwarewalkthrough": "toolmachinerysoftwarewalkthrough",
        "simulation": "simulationanimationmodelvisualisation",
        "simulationanimationmodelvisualisation": "simulationanimationmodelvisualisation",
    },
    "Dataset": {
        "entitydataset": "entityfocuseddatasetfarmfielddata",
        "entityfocuseddatasetfarmfielddata": "entityfocuseddatasetfarmfielddata",
        "operationsdataset": "eventoperationsdatasetactivityrecords",
        "eventoperationsdatasetactivityrecords": "eventoperationsdatasetactivityrecords",
        "inputusedataset": "inputusedatasetinputusedata",
        "inputusedatasetinputusedata": "inputusedatasetinputusedata",
        "outputdataset": "outputproductiondatasetproductiondata",
        "outputproductiondatasetproductiondata": "outputproductiondatasetproductiondata",
        "timeseriesdataset": "timeseriesdatasetweathertimedata",
        "timeseriesdatasetweathertimedata": "timeseriesdatasetweathertimedata",
        "geospatialdataset": "geospatialdatasetmapbasedorgeospatialdata",
        "geospatialdatasetmapbasedorgeospatialdata": "geospatialdatasetmapbasedorgeospatialdata",
        "analyticaldataset": "analyticalderiveddatasetanalysisinsightsdata",
        "analyticalderiveddatasetanalysisinsightsdata": "analyticalderiveddatasetanalysisinsightsdata",
        "surveydataset": "surveysocialdatasetfarmersurveydata",
        "surveysocialdatasetfarmersurveydata": "surveysocialdatasetfarmersurveydata",
        "machinedataset": "machineequipmentdatasetmachinerysensordata",
        "machineequipmentdatasetmachinerysensordata": "machineequipmentdatasetmachinerysensordata",
    },
    "Software": {
        "fmis": "farmmanagementsystemfmis",
        "farmmanagementsystemfmis": "farmmanagementsystemfmis",
        "monitoring": "monitoringrecordingtools",
        "monitoringrecordingtools": "monitoringrecordingtools",
        "fielddatacollection": "fielddatacollectionapps",
        "fielddatacollectionapps": "fielddatacollectionapps",
        "gismapping": "mappinggistools",
        "mappinggistools": "mappinggistools",
        "analysistools": "dataanalysisdashboardtools",
        "dataanalysisdashboardtools": "dataanalysisdashboardtools",
        "simulationtools": "simulationforecastingtools",
        "simulationforecastingtools": "simulationforecastingtools",
        "automationsystems": "automationcontrolsystems",
        "automationcontrolsystems": "automationcontrolsystems",
        "trainingapps": "traininglearningapplications",
        "traininglearningapplications": "traininglearningapplications",
    },
}


@dataclass
class ExtractedDoc:
    source_key: str
    source_path: Path
    category: str | None
    title_lines: list[str]
    tables: list[dict[str, Any]]


def slugify(text: str) -> str:
    text = text.lower().replace("&", " and ")
    text = re.sub(r"[^a-z0-9]+", "_", text)
    return text.strip("_")


def compact_spaces(text: str) -> str:
    return " ".join(text.split())


def split_semicolon_list(text: str) -> list[str]:
    return [item.strip() for item in text.split(";") if item.strip()]


def split_comma_list(text: str) -> list[str]:
    return [item.strip() for item in text.split(",") if item.strip()]


def extract_doc(spec_key: str, path: Path, category: str | None) -> ExtractedDoc:
    doc = Document(path)
    title_lines = [compact_spaces(p.text) for p in doc.paragraphs if compact_spaces(p.text)]
    tables: list[dict[str, Any]] = []
    for idx, table in enumerate(doc.tables):
        rows = [[compact_spaces(cell.text) for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        headers = rows[0]
        body = rows[1:]
        tables.append(
            {
                "table_index": idx,
                "headers": headers,
                "rows": [dict(zip(headers, row)) for row in body],
            }
        )
    return ExtractedDoc(
        source_key=spec_key,
        source_path=path,
        category=category,
        title_lines=title_lines,
        tables=tables,
    )


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def build_unified_subcategories(doc: ExtractedDoc) -> list[dict[str, Any]]:
    rows = doc.tables[0]["rows"]
    result = []
    for row in rows:
        subcategory = row["Subcategory"]
        subcategory_id = slugify(subcategory)
        result.append(
            {
                "id": subcategory_id,
                "name": subcategory,
                "definition": row["Definition"],
                "scope_note": row["Scope"],
                "user_label": row["User Label"],
                "detailed_features": split_semicolon_list(row["Detailed Features"]),
                "mapped_from_legacy": split_comma_list(row["Mapped From (Previous Subcategories)"]),
                "applicability_hints": UNIFIED_APPLICABILITY_HINTS.get(subcategory_id, []),
                "source_doc": doc.source_path.name,
            }
        )
    return result


def build_document_profiles(doc: ExtractedDoc) -> list[dict[str, Any]]:
    rows = doc.tables[0]["rows"]
    result = []
    for row in rows:
        type_name = row["Subcategory"]
        result.append(
            {
                "id": slugify(type_name),
                "name": type_name,
                "definition": row["Definition"],
                "scope_note": row["Scope"],
                "user_label": row["User Label"],
                "key_features": split_comma_list(row["Key Features"]),
                "merged_from": split_comma_list(row["Merged From"]),
                "category": doc.category,
                "source_doc": doc.source_path.name,
            }
        )
    return result


def normalize_rule_target(category: str, label: str) -> str:
    normalized = re.sub(r"[^a-z0-9]+", "", label.lower())
    aliases = CATEGORY_TYPE_ALIASES.get(category, {})
    return aliases.get(normalized, normalized)


def build_category_profile(doc: ExtractedDoc) -> dict[str, Any]:
    type_rows = doc.tables[0]["rows"]
    feature_rows = doc.tables[1]["rows"]
    rule_rows = doc.tables[2]["rows"]

    features = []
    for row in feature_rows:
        features.append(
            {
                "id": slugify(row["Feature"]),
                "name": row["Feature"],
                "definition": row["Definition"],
                "distinguishes": split_comma_list(row["Distinguishes"]),
            }
        )

    rule_by_type_id: dict[str, list[str]] = {}
    for row in rule_rows:
        key = normalize_rule_target(doc.category or "", row["Type"])
        rule_by_type_id.setdefault(key, []).append(row["Dominant rule"])

    types = []
    for row in type_rows:
        type_name = row["Type"]
        type_id = normalize_rule_target(doc.category or "", type_name)
        types.append(
            {
                "id": type_id,
                "name": type_name,
                "definition": row["Definition"],
                "scope_note": row["Scope note"],
                "examples": split_comma_list(row["Examples"]),
                "dominant_rules": rule_by_type_id.get(type_id, []),
                "source_doc": doc.source_path.name,
            }
        )

    return {
        "category": doc.category,
        "source_doc": doc.source_path.name,
        "types": types,
        "feature_catalog": features,
    }


def build_combined_model(extracted_docs: dict[str, ExtractedDoc]) -> dict[str, Any]:
    unified = build_unified_subcategories(extracted_docs["unified"])
    document_profiles = build_document_profiles(extracted_docs["document"])
    category_profiles = {
        key.capitalize() if key != "dataset" else "Dataset": build_category_profile(doc)
        for key, doc in extracted_docs.items()
        if key in {"image", "audio", "video", "dataset", "software"}
    }

    return {
        "schema_version": "v4",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "design_status": "source_model_only",
        "design_notes": {
            "subcategory_binding": "category_agnostic",
            "file_category_derivation": "deterministic_from_mime_or_extension",
            "url_category_derivation": "content_or_metadata_based",
            "note": (
                "Unified subcategories are reusable semantic concepts. "
                "Category-specific type profiles remain available as evidence models and signal families, "
                "not as hard subcategory parents."
            ),
        },
        "source_docs": {key: str(doc.source_path.relative_to(BASE_DIR)) for key, doc in extracted_docs.items()},
        "unified_subcategories": unified,
        "document_profiles": document_profiles,
        "category_type_profiles": category_profiles,
    }


def main() -> None:
    extracted_docs = {
        key: extract_doc(key, spec["path"], spec["category"])
        for key, spec in DOC_SPECS.items()
    }

    for key, doc in extracted_docs.items():
        raw_payload = {
            "source_key": doc.source_key,
            "source_path": str(doc.source_path.relative_to(BASE_DIR)),
            "category": doc.category,
            "title_lines": doc.title_lines,
            "tables": doc.tables,
        }
        write_json(OUT_DIR / f"{doc.source_path.stem}.json", raw_payload)

    combined = build_combined_model(extracted_docs)
    write_json(OUT_DIR / "subcategory_model_v4.json", combined)
    print(f"[OK] Wrote combined subtype model -> {OUT_DIR / 'subcategory_model_v4.json'}")


if __name__ == "__main__":
    main()
