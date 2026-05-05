from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document


BASE_DIR = Path(__file__).resolve().parents[1]
DATA_MODEL_DIR = BASE_DIR / "data_model"
NEW_DIR = DATA_MODEL_DIR / "new"
OUT_DIR = DATA_MODEL_DIR / "generated" / "v5"

MONGO_EXPORT_PATH = DATA_MODEL_DIR / "data_model.subcategories.v5.json"
FULL_MODEL_PATH = OUT_DIR / "subcategories_v5_full_model.json"

SOURCE_DOCS = {
    "document": NEW_DIR / "Document_types_revised_v4.docx",
    "audio": NEW_DIR / "Audio_types_and_features_v4.docx",
    "video": NEW_DIR / "Video_types_and_features_v4.docx",
    "image": NEW_DIR / "Image_types_and_features_v4.docx",
    "dataset": NEW_DIR / "Dataset_types_and_features_v4.docx",
    "software": NEW_DIR / "Software_types_and_features_v4.docx",
    "unified": NEW_DIR / "Unified_subcategories_mapping_v5_exhaustive.docx",
}

MONGO_CATEGORY_IDS = {
    "Document": "hkCpEjKW",
    "Video": "H4AQkKXG",
    "Audio": "53DSgGhE",
    "Image": "7oeWtUoi",
    "Dataset": "SdtDnWxr",
    "Software Application": "6NhmfRWh",
}

MODALITY_TO_CATEGORY = {
    "Document": "Document",
    "Video": "Video",
    "Audio": "Audio",
    "Image": "Image",
    "Dataset": "Dataset",
    "Software": "Software Application",
}

MODALITY_SOURCE_KEYS = {
    "Document": "document",
    "Audio": "audio",
    "Video": "video",
    "Image": "image",
    "Dataset": "dataset",
    "Software": "software",
}

MODALITY_TYPE_HEADERS = {
    "Document": "Subcategory",
    "Audio": "Type",
    "Video": "Type",
    "Image": "Type",
    "Dataset": "Type",
    "Software": "Type",
}

PROFILE_ALIASES = {
    "Audio": {
        "Interview/practitioner perspective": ["Interview"],
        "Expert commentary/explainer": ["Explainer"],
        "Case study/practice story": ["Case study"],
        "Panel discussion": ["Panel"],
        "Q&A": ["Q&A"],
        "Talk/lecture": ["Lecture"],
        "How-to/procedure guide": ["How-to"],
    },
    "Video": {
        "Field demonstration/walkthrough": ["Walkthrough"],
        "How-to/procedure demonstration": ["How-to"],
        "Case study": ["Case study"],
        "Explainer/documentary": ["Explainer"],
        "Interview/practitioner perspective": ["Interview"],
        "Expert Q&A session": ["Q&A"],
        "Panel discussion": ["Panel"],
        "Recorded presentation/webinar": ["Presentation"],
        "Tool/machinery/software walkthrough": ["Tool walkthrough"],
        "Simulation/animation/model visualisation": ["Simulation"],
    },
    "Image": {
        "Chart/graph": ["Chart/graph"],
        "Infographic": ["Infographic"],
        "Diagram/schematic": ["Diagram/schematic"],
        "Map": ["Map"],
        "Field/observational photograph": ["Field/observational photograph"],
        "Diagnostic photograph": ["Diagnostic photograph"],
        "Equipment/infrastructure photograph": ["Equipment/infrastructure photograph"],
        "Aerial/remote-sensing image": ["Aerial/remote-sensing image"],
    },
    "Dataset": {
        "Entity-focused dataset (Farm/Field Data)": ["Entity dataset"],
        "Event/operations dataset (Activity Records)": ["Operations dataset"],
        "Input-use dataset (Input Use Data)": ["Input-use dataset"],
        "Output/production dataset (Production Data)": ["Output dataset"],
        "Time-series dataset (Weather & Time Data)": ["Time-series dataset"],
        "Geospatial dataset (Map-based or Geospatial Data)": ["Geospatial dataset"],
        "Analytical/derived dataset (Analysis & Insights Data)": ["Analytical dataset"],
        "Survey/social dataset (Farmer & Survey Data)": ["Survey dataset"],
        "Machine/equipment dataset (Machinery & Sensor Data)": ["Machine dataset"],
    },
    "Software": {
        "Farm Management System (FMIS)": ["FMIS"],
        "Monitoring & Recording Tools": ["Monitoring"],
        "Field Data Collection Apps": ["Field data collection"],
        "Mapping & GIS Tools": ["GIS/Mapping"],
        "Data Analysis & Dashboard Tools": ["Analysis tools"],
        "Simulation & Forecasting Tools": ["Simulation tools"],
        "Automation & Control Systems": ["Automation systems"],
        "Training & Learning Applications": ["Training apps"],
    },
}


def compact(text: str) -> str:
    return " ".join((text or "").split())


def slugify(text: str) -> str:
    text = text.lower().replace("&", " and ")
    text = re.sub(r"[^a-z0-9]+", "_", text)
    return text.strip("_")


def normalize(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", (text or "").lower())


def split_semicolons(text: str) -> list[str]:
    return [compact(item) for item in (text or "").split(";") if compact(item)]


def split_commas(text: str) -> list[str]:
    return [compact(item) for item in (text or "").split(",") if compact(item)]


def deterministic_id(name: str) -> str:
    digest = hashlib.sha1(name.encode("utf-8")).hexdigest()
    return digest[:8]


def rows_from_table(table) -> list[list[str]]:
    return [[compact(cell.text) for cell in row.cells] for row in table.rows]


def dict_rows(table) -> list[dict[str, str]]:
    rows = rows_from_table(table)
    if not rows:
        return []
    headers = rows[0]
    body = rows[1:]
    return [dict(zip(headers, row)) for row in body]


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def parse_unified_doc(path: Path) -> dict[str, Any]:
    doc = Document(path)
    tables = doc.tables
    legend_rows = dict_rows(tables[0])
    summary_rows = dict_rows(tables[1])
    details_tables = tables[2:50]
    appendix_rows = dict_rows(tables[50])

    legend = {row["Strength"]: row["Meaning"] for row in legend_rows}
    appendix_index = {
        (row["Modality"], row["Source subcategory"]): row["Unified mapping(s)"]
        for row in appendix_rows
    }

    unified_items: list[dict[str, Any]] = []
    for index in range(0, len(details_tables), 2):
        detail_table = details_tables[index]
        mapping_table = details_tables[index + 1]
        detail_rows = dict_rows(detail_table)
        mapping_rows = dict_rows(mapping_table)
        detail_map = {row["Field"]: row["Value"] for row in detail_rows}
        summary_row = summary_rows[index // 2]
        name = summary_row["Unified subcategory"]
        item_id = slugify(name)
        mappings = []
        categories = []
        for row in mapping_rows:
            category = MODALITY_TO_CATEGORY[row["Modality"]]
            if category not in categories:
                categories.append(category)
            mappings.append(
                {
                    "modality": row["Modality"],
                    "category": category,
                    "source_subcategory": row["Source subcategory"],
                    "source_subcategory_id": slugify(row["Source subcategory"]),
                    "strength": row["Strength"],
                    "strength_definition": legend[row["Strength"]],
                    "rationale": row["Rationale"],
                }
            )

        unified_items.append(
            {
                "id": item_id,
                "name": name,
                "user_label": detail_map.get("User label", summary_row["User-facing label"]),
                "definition": detail_map.get("Definition", ""),
                "scope_note": detail_map.get("Scope note", ""),
                "feature_basis": split_semicolons(detail_map.get("Feature basis", "")),
                "summary_counts": {
                    "strong_mappings": int(summary_row["Strong mappings"]),
                    "partial_or_weak_mappings": int(summary_row["Partial/weak mappings"]),
                },
                "applicable_categories": categories,
                "source_mappings": mappings,
                "source_doc": path.name,
                "appendix_references": [
                    {
                        "modality": row["Modality"],
                        "source_subcategory": row["Source subcategory"],
                        "unified_mapping_text": appendix_index.get((row["Modality"], row["Source subcategory"]), ""),
                    }
                    for row in mapping_rows
                ],
            }
        )

    return {
        "title": compact(doc.paragraphs[0].text),
        "subtitle": compact(doc.paragraphs[1].text),
        "mapping_strength_legend": legend,
        "summary_table": summary_rows,
        "appendix_source_to_unified_mappings": appendix_rows,
        "unified_subcategories": unified_items,
    }


def parse_modality_doc(modality: str, path: Path) -> dict[str, Any]:
    doc = Document(path)
    title_lines = [compact(p.text) for p in doc.paragraphs if compact(p.text)]
    type_rows = dict_rows(doc.tables[0])
    feature_rows = dict_rows(doc.tables[1]) if len(doc.tables) > 1 else []
    rule_rows = dict_rows(doc.tables[2]) if len(doc.tables) > 2 else []

    feature_by_distinguishes: dict[str, list[dict[str, str]]] = {}
    for row in feature_rows:
        key = normalize(row["Distinguishes"])
        feature_by_distinguishes.setdefault(key, []).append(
            {
                "feature_id": row["Feature"],
                "definition": row["Definition"],
                "distinguishes": row["Distinguishes"],
            }
        )

    rule_by_type = {normalize(row["Type"]): row["Dominant rule"] for row in rule_rows}
    header_name = MODALITY_TYPE_HEADERS[modality]

    profiles = []
    for row in type_rows:
        type_name = row[header_name]
        preferred_aliases = [normalize(alias) for alias in PROFILE_ALIASES.get(modality, {}).get(type_name, [])]
        candidates = {
            normalize(type_name),
            normalize(type_name.split("(")[0]),
            normalize(type_name.replace("/", " ")),
        }
        for alias_key in preferred_aliases:
            candidates.add(alias_key)

        matched_features: list[dict[str, str]] = []
        seen_feature_ids: set[str] = set()
        for alias_key in preferred_aliases:
            for item in feature_by_distinguishes.get(alias_key, []):
                if item["feature_id"] not in seen_feature_ids:
                    matched_features.append(item)
                    seen_feature_ids.add(item["feature_id"])
        if not matched_features:
            for key, items in feature_by_distinguishes.items():
                for candidate in candidates:
                    if key and (key in candidate or candidate in key):
                        for item in items:
                            if item["feature_id"] not in seen_feature_ids:
                                matched_features.append(item)
                                seen_feature_ids.add(item["feature_id"])
                        break

        dominant_rule = ""
        for alias_key in preferred_aliases:
            if alias_key in rule_by_type:
                dominant_rule = rule_by_type[alias_key]
                break
        if not dominant_rule:
            for candidate in candidates:
                if candidate in rule_by_type:
                    dominant_rule = rule_by_type[candidate]
                    break
        if not dominant_rule:
            for key, value in rule_by_type.items():
                for candidate in candidates:
                    if key and (key in candidate or candidate in key):
                        dominant_rule = value
                        break
                if dominant_rule:
                    break

        profiles.append(
            {
                "id": slugify(type_name),
                "name": type_name,
                "modality": modality,
                "category": MODALITY_TO_CATEGORY[modality],
                "definition": row.get("Definition", ""),
                "scope_note": row.get("Scope note", row.get("Scope", "")),
                "examples": row.get("Examples", ""),
                "user_label": None,
                "key_features_text": None,
                "merged_from": [],
                "feature_catalog_matches": matched_features,
                "dominant_rule": dominant_rule,
                "source_doc": path.name,
            }
        )

    if modality == "Document":
        for profile, row in zip(profiles, type_rows):
            profile["user_label"] = row.get("User Label", "")
            profile["key_features_text"] = row.get("Key Features", "")
            profile["merged_from"] = split_commas(row.get("Merged From", ""))
            profile["feature_catalog_matches"] = [
                {
                    "feature_id": slugify(feature),
                    "definition": feature,
                    "distinguishes": profile["name"],
                }
                for feature in split_commas(row.get("Key Features", ""))
            ]
            profile["dominant_rule"] = f"Must align with: {row.get('Scope', '')}" if row.get("Scope") else ""

    return {
        "title_lines": title_lines,
        "profiles": profiles,
        "feature_catalog": feature_rows,
        "dominant_rules": rule_rows,
        "source_doc": path.name,
    }


def attach_cross_links(full_model: dict[str, Any]) -> None:
    source_profile_index: dict[tuple[str, str], dict[str, Any]] = {}
    for modality, payload in full_model["source_modalities"].items():
        for profile in payload["profiles"]:
            source_profile_index[(modality, profile["name"])] = profile

    for unified in full_model["unified_subcategories"]:
        unified["source_profiles"] = []
        for mapping in unified["source_mappings"]:
            profile = source_profile_index.get((mapping["modality"], mapping["source_subcategory"]))
            if not profile:
                continue
            unified["source_profiles"].append(
                {
                    "modality": mapping["modality"],
                    "category": mapping["category"],
                    "profile_id": profile["id"],
                    "profile_name": profile["name"],
                    "strength": mapping["strength"],
                    "rationale": mapping["rationale"],
                    "definition": profile["definition"],
                    "scope_note": profile["scope_note"],
                    "dominant_rule": profile["dominant_rule"],
                    "feature_catalog_matches": profile["feature_catalog_matches"],
                    "examples": profile["examples"],
                    "merged_from": profile["merged_from"],
                }
            )

    unified_by_name = {item["name"]: item for item in full_model["unified_subcategories"]}
    appendix_rows = full_model["appendix_source_to_unified_mappings"]
    for row in appendix_rows:
        modality = row["Modality"]
        source_name = row["Source subcategory"]
        profile = source_profile_index.get((modality, source_name))
        if not profile:
            continue
        parsed_targets = []
        for piece in split_semicolons(row["Unified mapping(s)"]):
            match = re.match(r"^(.*?)\s+\((Strong|Partial|Weak/Partial)\)$", piece)
            if not match:
                continue
            target_name, strength = match.groups()
            target = unified_by_name.get(target_name)
            parsed_targets.append(
                {
                    "unified_subcategory_id": target["id"] if target else slugify(target_name),
                    "unified_subcategory_name": target_name,
                    "strength": strength,
                }
            )
        profile["unified_mappings"] = parsed_targets


def build_mongo_export(full_model: dict[str, Any]) -> list[dict[str, Any]]:
    now = datetime.now(timezone.utc).isoformat(timespec="milliseconds").replace("+00:00", "Z")
    records = []
    for item in full_model["unified_subcategories"]:
        categories = item["applicable_categories"]
        category_ids = [MONGO_CATEGORY_IDS[name] for name in categories]
        source_modalities = sorted({mapping["modality"] for mapping in item["source_mappings"]})
        source_subcategories = sorted(
            {
                mapping["source_subcategory"]
                for mapping in item["source_mappings"]
            }
        )
        records.append(
            {
                "_id": deterministic_id(item["id"]),
                "created_ts": {"$date": now},
                "created_by": "System",
                "updated_ts": {"$date": now},
                "updated_by": "System",
                "version": 1,
                "status": "Published",
                "name": item["name"],
                "user_label": item["user_label"],
                "definition": item["definition"],
                "scope_note": item["scope_note"],
                "feature_basis": item["feature_basis"],
                "parent_category": categories,
                "parent_category_id": category_ids,
                "category_binding_mode": "decoupled",
                "subcategory_type": "unified",
                "source_modalities": source_modalities,
                "source_subcategories": source_subcategories,
                "mapping_strength_counts": item["summary_counts"],
                "source_mappings": item["source_mappings"],
                "source_profiles": item["source_profiles"],
                "source_docs": sorted(
                    {
                        item["source_doc"],
                        *[
                            profile["modality"] + ":" + profile["profile_name"]
                            for profile in item["source_profiles"]
                        ],
                    }
                ),
            }
        )
    return records


def main() -> None:
    unified_payload = parse_unified_doc(SOURCE_DOCS["unified"])
    source_modalities = {
        modality: parse_modality_doc(modality, path)
        for modality, path in (
            ("Document", SOURCE_DOCS["document"]),
            ("Audio", SOURCE_DOCS["audio"]),
            ("Video", SOURCE_DOCS["video"]),
            ("Image", SOURCE_DOCS["image"]),
            ("Dataset", SOURCE_DOCS["dataset"]),
            ("Software", SOURCE_DOCS["software"]),
        )
    }

    full_model = {
        "model_version": "v5",
        "created_from": {key: str(path.relative_to(BASE_DIR)) for key, path in SOURCE_DOCS.items()},
        "category_policy": {
            "category_is_derived_from_mimetype": True,
            "subcategory_is_signal_derived": True,
            "category_and_subcategory_are_decoupled": True,
        },
        "mapping_strength_legend": unified_payload["mapping_strength_legend"],
        "unified_subcategories": unified_payload["unified_subcategories"],
        "source_modalities": source_modalities,
        "appendix_source_to_unified_mappings": unified_payload["appendix_source_to_unified_mappings"],
        "summary_table": unified_payload["summary_table"],
        "source_titles": {
            key: payload["title_lines"] if isinstance(payload, dict) and "title_lines" in payload else []
            for key, payload in source_modalities.items()
        },
    }
    attach_cross_links(full_model)

    mongo_export = build_mongo_export(full_model)
    write_json(FULL_MODEL_PATH, full_model)
    write_json(MONGO_EXPORT_PATH, mongo_export)
    print(f"[OK] Wrote {FULL_MODEL_PATH}")
    print(f"[OK] Wrote {MONGO_EXPORT_PATH}")


if __name__ == "__main__":
    main()
