from __future__ import annotations

from docx import Document

from docint.ingest.models import IngestedAsset


def ingest_docx(file_path: str, filename: str) -> IngestedAsset:
    doc = Document(file_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_values:
                lines.append(" | ".join(row_values))

    text = "\n".join(lines).strip()

    return IngestedAsset(
        asset_type="docx",
        filename=filename,
        source_path=file_path,
        text=text,
        lines=lines,
        units=max(1, len(doc.sections)),
        unit_label="sections",
        source="docx_text",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        visual_candidate=False,
        ocr_supported=False,
        meta={
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        },
    )
